"""
Tải và chunk tài liệu bằng LangChain — bổ sung (không thay) extract/split legacy trong ingest_utils.
Video / OCR vẫn dùng ingest_utils.extract_text khi cần.
"""

from __future__ import annotations

import json
import os
from pathlib import Path
from typing import List

from langchain_core.documents import Document

SUPPORTED_LC_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".html", ".htm", ".csv", ".json"}


def _single_doc(text: str, source_stem: str, file_path: str) -> List[Document]:
    t = (text or "").strip()
    if not t:
        return []
    return [Document(page_content=t, metadata={"source": source_stem, "file_path": file_path})]


def _load_docx_python(path: Path) -> List[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    return _single_doc(text, path.stem, str(path))


def load_document(file_path: str) -> List[Document]:
    """
    Load file thành list Document. Luôn gắn metadata source (stem) + file_path.
    Fallback về ingest_utils.extract_text nếu loader LangChain thất bại.
    """
    path = Path(file_path)
    if not path.is_file():
        return []
    ext = path.suffix.lower()
    stem = path.stem
    fp = str(path)

    if ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".doc") or ext not in SUPPORTED_LC_EXTENSIONS:
        from app.domains.ingest.ingest_utils import extract_text
        return _single_doc(extract_text(fp), stem, fp)

    try:
        if ext == ".pdf":
            from langchain_community.document_loaders import PyMuPDFLoader

            return PyMuPDFLoader(fp).load()

        if ext in (".txt", ".md"):
            from langchain_community.document_loaders import TextLoader

            return TextLoader(fp, encoding="utf-8", autodetect_encoding=True).load()

        if ext == ".docx":
            return _load_docx_python(path)

        if ext in (".html", ".htm"):
            from langchain_community.document_loaders import UnstructuredHTMLLoader

            return UnstructuredHTMLLoader(fp).load()

        if ext == ".csv":
            from langchain_community.document_loaders import CSVLoader

            return CSVLoader(fp).load()

        if ext == ".json":
            raw = path.read_text(encoding="utf-8", errors="ignore")
            try:
                obj = json.loads(raw)
                text = json.dumps(obj, ensure_ascii=False, indent=2) if not isinstance(obj, str) else obj
            except Exception:
                text = raw
            return _single_doc(text, stem, fp)

    except Exception:
        pass

    from app.domains.ingest.ingest_utils import extract_text
    return _single_doc(extract_text(fp), stem, fp)


def split_documents(
    docs: List[Document],
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> List[Document]:
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    cs = chunk_size if chunk_size is not None else int(os.getenv("CHUNK_SIZE", "500"))
    co = chunk_overlap if chunk_overlap is not None else int(os.getenv("CHUNK_OVERLAP", "50"))
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=cs,
        chunk_overlap=co,
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
        length_function=len,
    )
    out = splitter.split_documents(docs)
    for i, d in enumerate(out):
        d.metadata.setdefault("chunk_index", i)
        d.metadata["total_chunks"] = len(out)
    return out
