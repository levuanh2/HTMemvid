import os
import fitz  # PyMuPDF để đọc PDF
from docx import Document
from typing import List
import pytesseract
from PIL import Image
import shutil  # Để tìm tesseract trong PATH
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_experimental.text_splitter import SemanticChunker
# Auto-detect Tesseract path (portable cho Windows/Linux/Mac)
tesseract_path = shutil.which('tesseract') or r"C:\Program Files\Tesseract-OCR\tesseract.exe"
if tesseract_path and os.path.exists(tesseract_path):
    pytesseract.pytesseract.tesseract_cmd = tesseract_path
else:
    print("⚠️ Tesseract not found - OCR will fail for images.")
    
FRAME_DIR = "qr_frames"
os.makedirs(FRAME_DIR, exist_ok=True)

def extract_text(path: str) -> str:
    ext = path.lower()

    # PDF
    if ext.endswith('.pdf'):
        try:
            doc = fitz.open(path)
            text = ''.join(page.get_text() for page in doc)
            doc.close()
            return text
        except:
            return ''

    # DOCX
    if ext.endswith('.docx'):
        try:
            doc = Document(path)
            return '\n'.join(p.text for p in doc.paragraphs)
        except:
            return ''

    # DOC (Word 97-2003)
    if ext.endswith('.doc'):
        try:
            import subprocess
            # Chuyển sang docx bằng libreoffice nếu có
            temp_docx = path + ".docx"
            subprocess.run(["soffice", "--headless", "--convert-to", "docx", path, "--outdir", os.path.dirname(path)])
            doc = Document(temp_docx)
            return '\n'.join(p.text for p in doc.paragraphs)
        except:
            return ''

    # TXT
    if ext.endswith('.txt'):
        try:
            return open(path, encoding='utf-8', errors='ignore').read()
        except:
            return ''

    # Hình ảnh (OCR)
    if ext.endswith(('.png', '.jpg', '.jpeg')):
        try:
            img = Image.open(path)
            text = pytesseract.image_to_string(img, lang='eng+vie')
            return text
        except:
            return ''

    return ''


import re 
_semantic_embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
def split_text(text: str) -> List[str]:
    """
    Chia van ban thanbh cac chunk dua tren ngu nghia (semantic chunk)
    su dung cung model embedding vs FAISS
    """
    if not text or not text.strip():
        return []
    #tao semantic
    text_splitter=SemanticChunker(
        _semantic_embeddings,
        breakpoint_threshold_type="percentile",
        breakpoint_threshold_amount=94,
        min_chunk_size=300,
    )
    chunks=text_splitter.split_text(text.strip())
    chunks= [chunk.strip() for chunk in chunks if chunk.strip()]
    return chunks
# def split_text(text: str, max_bytes: int = 1050, overlap_ratio: float = 0.2):
#     # Tách văn bản thành các câu
#
#     sentences = re.split(r'(?<=[.!?])\s+', text)
#     chunks, cur = [], []
#     cur_bytes = 0
#
#     for sent in sentences:
#         if not sent.strip():
#             continue
#         sent_bytes = len(sent.encode('utf-8'))
#         if cur_bytes + sent_bytes > max_bytes:
#             if cur:
#                 chunk_text = " ".join(cur).strip()
#                 chunks.append(chunk_text)
#                 # lấy overlap = 20% câu cuối của chunk
#                 overlap_size = max(1, int(len(cur) * overlap_ratio))
#                 cur = cur[-overlap_size:]
#                 cur_bytes = sum(len(s.encode('utf-8')) for s in cur)
#         cur.append(sent)
#         cur_bytes += sent_bytes
#
#     if cur:
#         chunks.append(" ".join(cur).strip())
#
#     return chunks
