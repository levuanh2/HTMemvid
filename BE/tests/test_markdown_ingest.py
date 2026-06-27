from pathlib import Path

from docx import Document

from app.domains.ingest.clean import clean_markdown
from app.domains.ingest.markdown_convert import convert_and_save, to_markdown


def test_to_markdown_docx_emits_heading(tmp_path):
    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("Test Heading", level=1)
    document.add_paragraph("Doan van ban thu nhat.")
    document.add_paragraph("Doan van ban thu hai.")
    document.save(docx_path)

    markdown = to_markdown(str(docx_path))

    assert "#" in markdown


def test_clean_markdown_removes_repeated_footer_and_dehyphenates():
    raw = "\n".join(
        [
            "Noi dung mo dau",
            "Tai lieu noi bo",
            "word-",
            "break",
            "1",
            "Noi dung tiep theo",
            "Tai lieu noi bo",
            "2",
            "Ket thuc",
            "Tai lieu noi bo",
            "3",
        ]
    )

    cleaned = clean_markdown(raw)

    assert "Tai lieu noi bo" not in cleaned
    assert "wordbreak" in cleaned
    assert "1\n" not in cleaned


def test_convert_and_save_writes_markdown_file(tmp_path):
    source_path = tmp_path / "note.txt"
    source_path.write_text("Xin chao markdown", encoding="utf-8")

    markdown, saved_path = convert_and_save(str(source_path), md_dir=str(tmp_path / "md"))

    assert markdown == "Xin chao markdown"
    assert Path(saved_path).exists()
    assert Path(saved_path).suffix == ".md"
